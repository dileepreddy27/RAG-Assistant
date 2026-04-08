from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Any

from fastapi import UploadFile
from pypdf import PdfReader


@dataclass(slots=True)
class ParsedDocument:
    filename: str
    text: str
    metadata: dict[str, Any]


SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx"}


async def parse_upload_file(upload: UploadFile) -> ParsedDocument:
    filename = upload.filename or "uploaded_document"
    extension = Path(filename).suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type '{extension}'.")

    data = await upload.read()
    text = _extract_text(extension, data)

    cleaned = text.strip()
    if not cleaned:
        raise ValueError("No extractable text found in file.")

    return ParsedDocument(
        filename=filename,
        text=cleaned,
        metadata={
            "source_filename": filename,
            "extension": extension,
        },
    )


def _extract_text(extension: str, data: bytes) -> str:
    if extension == ".pdf":
        return _extract_pdf_text(data)

    if extension in {".txt", ".md"}:
        return data.decode("utf-8", errors="ignore")

    if extension == ".docx":
        from docx import Document as DocxDocument

        document = DocxDocument(BytesIO(data))
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    raise ValueError("Unsupported file type.")


def _extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)
